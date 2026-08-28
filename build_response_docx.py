#!/usr/bin/env python3
"""Build a polished response-to-reviewers DOCX from verified Markdown."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
NAVY = "17365D"
BLUE = "2F75B5"
PALE_BLUE = "DDEBF7"
PALE_RED = "FCE4D6"
GRAY = "666666"


def shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    element = properties.find(qn("w:shd"))
    if element is None:
        element = OxmlElement("w:shd")
        properties.append(element)
    element.set(qn("w:fill"), fill)


def set_repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def add_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(begin)
    paragraph.add_run()._r.append(text)
    paragraph.add_run()._r.append(end)


def add_runs(paragraph, text: str, color: str | None = None) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
        else:
            run = paragraph.add_run(part)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def configure(document: Document) -> None:
    section = document.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin, section.bottom_margin = Cm(1.7), Cm(1.7)
    section.left_margin, section.right_margin = Cm(2.0), Cm(2.0)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in (("Title", 22, NAVY), ("Heading 1", 14, NAVY), ("Heading 2", 11.5, BLUE)):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("GENO-2026-1699   |   MAJOR REVISION RESPONSE")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Engineering Optimization  •  ").font.size = Pt(8)
    add_field(footer, "PAGE")
    footer.add_run(" / ")
    add_field(footer, "NUMPAGES")


def masthead(document: Document, draft: bool) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(11.8)
    table.columns[1].width = Cm(4.7)
    left, right = table.rows[0].cells
    shade(left, NAVY)
    shade(right, "C00000" if draft else BLUE)
    left.vertical_alignment = right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("ENGINEERING OPTIMIZATION")
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(10)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("WORKING DRAFT" if draft else "READY FOR AUTHOR CHECK")
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(8)


def add_markdown(document: Document, text: str) -> None:
    lines = text.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("|-"):
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and lines[index + 1].startswith("|---"):
            rows = []
            rows.append([item.strip() for item in line.strip("|").split("|")])
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                rows.append([item.strip() for item in lines[index].strip("|").split("|")])
                index += 1
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for row_index, values in enumerate(rows):
                for column_index, value in enumerate(values):
                    cell = table.cell(row_index, column_index)
                    cell.text = ""
                    add_runs(cell.paragraphs[0], value)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    if row_index == 0:
                        shade(cell, NAVY)
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.size = Pt(8)
                    else:
                        for run in cell.paragraphs[0].runs:
                            run.font.size = Pt(8)
                if row_index == 0:
                    set_repeat_header(table.rows[row_index])
            continue
        if line.startswith("# "):
            p = document.add_paragraph(style="Title")
            add_runs(p, line[2:])
        elif line.startswith("## "):
            p = document.add_paragraph(style="Heading 1")
            add_runs(p, line[3:])
        elif line.startswith("### "):
            p = document.add_paragraph(style="Heading 2")
            add_runs(p, line[4:])
        elif line.startswith("> "):
            table = document.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            shade(cell, PALE_RED if "Author check" in line else PALE_BLUE)
            p = cell.paragraphs[0]
            add_runs(p, line[2:], "9C0006" if "Author check" in line else NAVY)
        elif re.match(r"^\d+\. ", line):
            p = document.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            add_runs(p, line[2:])
        else:
            p = document.add_paragraph()
            add_runs(p, line, "C00000" if "{{" in line or "REPLACE_WITH" in line else None)
        index += 1


def main() -> None:
    completed = ROOT / "docs" / "RESPONSE_TO_AE_COMPLETED.md"
    source = completed if completed.exists() else ROOT / "docs" / "RESPONSE_TO_AE_TEMPLATE.md"
    draft = source.name.endswith("TEMPLATE.md")
    output = ROOT / "docs" / (
        "Response_to_AE_COMPLETED.docx" if not draft else "Response_to_AE_WORKING_DRAFT.docx"
    )
    document = Document()
    configure(document)
    masthead(document, draft)
    add_markdown(document, source.read_text(encoding="utf-8"))
    document.core_properties.title = "Response to the Associate Editor — GENO-2026-1699"
    document.core_properties.subject = "Major revision and resubmission"
    document.core_properties.author = "Ku Junhua et al."
    document.save(output)
    print(output)


if __name__ == "__main__":
    main()
